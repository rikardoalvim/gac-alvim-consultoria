# -*- coding: utf-8 -*-
# Gerador de Parecer + Histórico + Pipeline + IA
# Autor: Alvim Consultoria :)

import os
import io
import csv
from datetime import datetime

import pandas as pd
import streamlit as st

from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph, Spacer, SimpleDocTemplate
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors

from PyPDF2 import PdfMerger, PdfReader
from docx import Document

from openai import OpenAI  # requer OPENAI_API_KEY configurada


# ==========================================================
# CONFIGURAÇÕES BÁSICAS
# ==========================================================

# Pasta padrão para logs e pareceres
BASE_DIR = r"C:\DOCS\PARECER"
os.makedirs(BASE_DIR, exist_ok=True)

LOG_PAR = os.path.join(BASE_DIR, "pareceres_log.csv")


# ==========================================================
# OPENAI / IA
# ==========================================================

def get_openai_client():
    """Retorna cliente OpenAI usando a variável de ambiente OPENAI_API_KEY."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


def gerar_campos_via_openai(texto_base: str):
    """
    Usa IA para gerar Nome, Resumo Profissional, Análise de Perfil e Conclusão
    a partir de texto (currículo + observações).
    """
    client = get_openai_client()
    if not client:
        raise RuntimeError("OPENAI_API_KEY não configurada no sistema.")

    prompt = f"""
Você é um assistente de recrutamento.
A partir do texto abaixo (currículo + observações), gere:

[NOME]
Nome completo do candidato

[RESUMO]
Resumo Profissional (em terceira pessoa, de forma objetiva)

[ANALISE]
Análise de Perfil (postura, comunicação, senioridade, aderência à vaga, potencial)

[CONCLUSAO]
Conclusão (síntese da indicação ou não do candidato, com tom profissional)

Texto do candidato:
\"\"\"{texto_base}\"\"\"
"""

    resp = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt,
    )

    text = resp.output[0].content[0].text

    nome = resumo = analise = concl = ""
    bloco = None

    for linha in text.splitlines():
        t = linha.strip()
        if t == "[NOME]":
            bloco = "nome"
            continue
        if t == "[RESUMO]":
            bloco = "resumo"
            continue
        if t == "[ANALISE]":
            bloco = "analise"
            continue
        if t == "[CONCLUSAO]":
            bloco = "concl"
            continue

        if bloco == "nome":
            nome += linha + "\n"
        elif bloco == "resumo":
            resumo += linha + "\n"
        elif bloco == "analise":
            analise += linha + "\n"
        elif bloco == "concl":
            concl += linha + "\n"

    return nome.strip(), resumo.strip(), analise.strip(), concl.strip()


# ==========================================================
# FUNÇÕES AUXILIARES (PDF / DOCX / PDF MERGE / PDF TEXT)
# ==========================================================

def _normalizar_linkedin(link):
    if not link or not link.strip():
        return None
    link = link.strip()
    if not (link.startswith("http://") or link.startswith("https://")):
        return "https://" + link
    return link


def extract_text_from_pdf(upload):
    """Extrai texto simples de um PDF."""
    if not upload:
        return ""
    reader = PdfReader(upload)
    text = ""
    for p in reader.pages:
        try:
            text += p.extract_text() + "\n"
        except Exception:
            pass
    return text


def build_parecer_pdf_to_bytes(
    empresa, cargo, nome, localidade, idade, pretensao,
    resumo_profissional, analise_perfil, conclusao_texto,
    linkedin=""
):
    buffer = io.BytesIO()

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontSize=14,
        spaceAfter=20,
    )
    header_style = ParagraphStyle(
        "Header",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=11,
        textColor=colors.HexColor("#003366"),
    )
    section_title = ParagraphStyle(
        "SecTitle",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#003366"),
        spaceBefore=12,
        spaceAfter=4,
    )
    text_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10.5,
        alignment=TA_JUSTIFY,
        leading=15,
    )

    def to_html(txt: str) -> str:
        return txt.replace("\n", "<br/>")

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=40,
        bottomMargin=40,
        leftMargin=50,
        rightMargin=50,
    )

    content = []
    # Cabeçalho
    content.append(Paragraph(f"<b>Empresa:</b> {empresa}", header_style))
    content.append(Paragraph(f"<b>Cargo:</b> {cargo}", header_style))
    content.append(Spacer(1, 20))
    content.append(Paragraph("🧩 PARECER DE TRIAGEM – CANDIDATO", title_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph(
        f"<b>Nome:</b> {nome}<br/><b>Localidade:</b> {localidade}",
        text_style
    ))

    linkedin_url = _normalizar_linkedin(linkedin)
    if linkedin_url:
        content.append(Paragraph(
            f'<b>LinkedIn:</b> <link href="{linkedin_url}">{linkedin_url}</link>',
            text_style
        ))

    # Seções
    content.append(Paragraph("🧾 <b>Resumo Profissional</b>", section_title))
    content.append(Paragraph(to_html(resumo_profissional), text_style))

    content.append(Paragraph("💡 <b>Análise de Perfil</b>", section_title))
    content.append(Paragraph(to_html(analise_perfil), text_style))

    content.append(Paragraph("➡️ <b>Conclusão</b>", section_title))
    content.append(Paragraph(to_html(conclusao_texto), text_style))

    content.append(Paragraph("💰 <b>Informações de Remuneração</b>", section_title))
    content.append(Paragraph(
        f"Idade: {idade}<br/>Pretensão Salarial: {pretensao}",
        text_style
    ))

    doc.build(content)
    return buffer.getvalue()


def build_parecer_docx_to_bytes(
    empresa, cargo, nome, localidade, idade, pretensao,
    resumo_profissional, analise_perfil, conclusao_texto,
    linkedin=""
):
    doc = Document()

    doc.add_heading("PARECER DE TRIAGEM – CANDIDATO", level=1)
    doc.add_paragraph(f"Empresa: {empresa}")
    doc.add_paragraph(f"Cargo: {cargo}")
    doc.add_paragraph(f"Nome: {nome}")
    doc.add_paragraph(f"Localidade: {localidade}")

    linkedin_url = _normalizar_linkedin(linkedin)
    if linkedin_url:
        doc.add_paragraph(f"LinkedIn: {linkedin_url}")

    doc.add_heading("Resumo Profissional", level=2)
    for bloco in resumo_profissional.split("\n\n"):
        doc.add_paragraph(bloco)

    doc.add_heading("Análise de Perfil", level=2)
    for bloco in analise_perfil.split("\n\n"):
        doc.add_paragraph(bloco)

    doc.add_heading("Conclusão", level=2)
    for bloco in conclusao_texto.split("\n\n"):
        doc.add_paragraph(bloco)

    doc.add_heading("Informações de Remuneração", level=2)
    doc.add_paragraph(f"Idade: {idade}")
    doc.add_paragraph(f"Pretensão Salarial: {pretensao}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def merge_pdfs_bytes(parecer_bytes: bytes, resume_path: str) -> bytes:
    """Mescla o parecer com o currículo (PDF) e retorna bytes do PDF final."""
    if not os.path.exists(resume_path):
        raise FileNotFoundError(f"Currículo não encontrado: {resume_path}")
    merger = PdfMerger()
    merger.append(io.BytesIO(parecer_bytes))
    merger.append(resume_path)

    out = io.BytesIO()
    merger.write(out)
    merger.close()
    return out.getvalue()


# ==========================================================
# FUNÇÕES DE LOG (HISTÓRICO / PIPELINE)
# ==========================================================

def registrar_parecer_log(
    data_hora,
    empresa,
    cargo,
    nome,
    localidade,
    idade,
    pretensao,
    linkedin,
    resumo_profissional,
    analise_perfil,
    conclusao_texto,
    formato,
    caminho_arquivo,
    status_etapa="Em avaliação",
    status_contratacao="Pendente",
    motivo_decline="",
):
    """Registra um parecer gerado em um arquivo CSV com TODOS os campos do formulário."""
    existe = os.path.exists(LOG_PAR)
    with open(LOG_PAR, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f, delimiter=";")
        if not existe:
            writer.writerow([
                "data_hora",
                "empresa",
                "cargo",
                "nome",
                "localidade",
                "idade",
                "pretensao",
                "linkedin",
                "resumo_profissional",
                "analise_perfil",
                "conclusao_texto",
                "formato",
                "caminho_arquivo",
                "status_etapa",
                "status_contratacao",
                "motivo_decline",
            ])
        writer.writerow([
            data_hora,
            empresa,
            cargo,
            nome,
            localidade,
            idade,
            pretensao,
            linkedin,
            resumo_profissional,
            analise_perfil,
            conclusao_texto,
            formato,
            caminho_arquivo,
            status_etapa,
            status_contratacao,
            motivo_decline,
        ])


def carregar_pareceres_log() -> pd.DataFrame:
    """Lê o CSV de pareceres e retorna um DataFrame (ou DF vazio se não existir)."""
    colunas_necessarias = [
        "data_hora",
        "empresa",
        "cargo",
        "nome",
        "localidade",
        "idade",
        "pretensao",
        "linkedin",
        "resumo_profissional",
        "analise_perfil",
        "conclusao_texto",
        "formato",
        "caminho_arquivo",
        "status_etapa",
        "status_contratacao",
        "motivo_decline",
    ]
    if not os.path.exists(LOG_PAR):
        return pd.DataFrame(columns=colunas_necessarias)

    df = pd.read_csv(LOG_PAR, sep=";", encoding="utf-8")
    for col in colunas_necessarias:
        if col not in df.columns:
            df[col] = ""
    return df


# ==========================================================
# UI STREAMLIT
# ==========================================================

st.set_page_config(page_title="Gerador de Parecer", page_icon="🧩", layout="wide")
st.title("🧩 Gerador de Parecer - Alvim Consultoria")
st.markdown(
    "App para gerar pareceres, registrar histórico e gerenciar pipeline de candidatos, "
    "com apoio opcional de IA (OpenAI)."
)

# ----------------------------------------------------------
# Inicializar session_state para os campos do formulário
# ----------------------------------------------------------

for campo, valor_padrao in {
    "empresa": "Smartcitizen",
    "cargo": "Analista Adm/Financeiro",
    "nome": "",
    "localidade": "Curitiba – PR",
    "idade": "",
    "pretensao": "",
    "linkedin": "",
    "resumo_profissional": "",
    "analise_perfil": "",
    "conclusao_texto": "",
}.items():
    if campo not in st.session_state:
        st.session_state[campo] = valor_padrao


# ----------------------------------------------------------
# Bloco IA: testar conexão + gerar a partir de PDF/obs
# ----------------------------------------------------------

with st.expander("🤖 IA - Geração automática (opcional)", expanded=False):
    col_ia1, col_ia2 = st.columns(2)
    with col_ia1:
        if st.button("🔌 Testar conexão com OpenAI"):
            client = get_openai_client()
            if not client:
                st.error("OPENAI_API_KEY não configurada no sistema.")
            else:
                try:
                    _ = client.responses.create(model="gpt-4.1-mini", input="OK?")
                    st.success("Conexão ok com a OpenAI (gpt-4.1-mini).")
                except Exception as e:
                    st.error(f"Erro ao conectar: {e}")
    st.write("Use o PDF do currículo + observações para gerar automaticamente Nome, Resumo, Análise e Conclusão.")

    uploaded_pdf = st.file_uploader("📎 Anexar PDF do currículo para IA", type=["pdf"])
    obs_ia = st.text_area("Observações complementares (opcional)", height=100)

    if st.button("✨ Gerar campos automaticamente via IA"):
        if not uploaded_pdf and not obs_ia.strip():
            st.warning("Anexe um PDF ou escreva observações para usar a IA.")
        else:
            try:
                texto_pdf = extract_text_from_pdf(uploaded_pdf) if uploaded_pdf else ""
                texto_base = (texto_pdf or "") + "\n\n" + (obs_ia or "")

                nome_ai, resumo_ai, analise_ai, concl_ai = gerar_campos_via_openai(texto_base)

                if nome_ai:
                    st.session_state["nome"] = nome_ai
                if resumo_ai:
                    st.session_state["resumo_profissional"] = resumo_ai
                if analise_ai:
                    st.session_state["analise_perfil"] = analise_ai
                if concl_ai:
                    st.session_state["conclusao_texto"] = concl_ai

                st.success("Campos preenchidos com sucesso. Role até o formulário para revisar/ajustar.")
            except Exception as e:
                st.error(f"Erro ao gerar via IA: {e}")


st.markdown("---")

# ----------------------------------------------------------
# FORMULÁRIO DE PARECER
# ----------------------------------------------------------

st.header("📝 Registro de novo parecer")

colA, colB = st.columns(2)
with colA:
    st.text_input("Empresa", key="empresa")
    st.text_input("Nome do candidato", key="nome")
    st.text_input("Idade", key="idade")
with colB:
    st.text_input("Cargo", key="cargo")
    st.text_input("Localidade", key="localidade")
    st.text_input("Pretensão Salarial (ex.: R$ 4.500,00)", key="pretensao")

st.text_input("LinkedIn (opcional)", key="linkedin")

st.subheader("Resumo profissional")
st.text_area(
    "Resumo profissional",
    height=150,
    key="resumo_profissional",
    placeholder="Descreva de forma objetiva o perfil profissional do candidato...",
)

st.subheader("Análise de perfil")
st.text_area(
    "Análise de perfil",
    height=150,
    key="analise_perfil",
    placeholder="Observações sobre postura, comunicação, senioridade, aderência à vaga...",
)

st.subheader("Conclusão")
st.text_area(
    "Conclusão",
    height=120,
    key="conclusao_texto",
    placeholder="Conclusão geral sobre o candidato e recomendação...",
)

st.markdown("---")

st.subheader("Arquivos e saída")

formato = st.radio("Formato do parecer", ["PDF", "DOCX"], index=0)

pasta_cv = st.text_input("Pasta com currículos PDF", value=BASE_DIR)
lista_cv = []
if os.path.isdir(pasta_cv):
    lista_cv = [f for f in os.listdir(pasta_cv) if f.lower().endswith(".pdf")]
selected_cv = st.selectbox(
    "Anexar currículo (PDF) ao parecer (opcional)",
    ["(Não anexar)"] + lista_cv,
)

output_folder = st.text_input("Pasta de saída dos pareceres", value=BASE_DIR)

nome_para_base = st.session_state["nome"] if st.session_state["nome"] else "Candidato"
nome_base = f"Parecer_{nome_para_base.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}"

if st.button("💾 Gerar parecer e registrar histórico"):
    nome = st.session_state["nome"]
    if not nome.strip():
        st.error("Informe o nome do candidato antes de gerar o parecer.")
    else:
        try:
            os.makedirs(output_folder, exist_ok=True)
            data_hora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            empresa = st.session_state["empresa"]
            cargo = st.session_state["cargo"]
            localidade = st.session_state["localidade"]
            idade = st.session_state["idade"]
            pretensao = st.session_state["pretensao"]
            linkedin = st.session_state["linkedin"]
            resumo_prof = st.session_state["resumo_profissional"]
            analise_prof = st.session_state["analise_perfil"]
            conclusao_txt = st.session_state["conclusao_texto"]

            if formato == "PDF":
                parecer_bytes = build_parecer_pdf_to_bytes(
                    empresa, cargo, nome, localidade, idade, pretensao,
                    resumo_prof, analise_prof, conclusao_txt,
                    linkedin
                )

                if selected_cv != "(Não anexar)":
                    resume_path = os.path.join(pasta_cv, selected_cv)
                    parecer_bytes = merge_pdfs_bytes(parecer_bytes, resume_path)

                filename = nome_base + ".pdf"
                caminho_final = os.path.join(output_folder, filename)

                with open(caminho_final, "wb") as f:
                    f.write(parecer_bytes)

                registrar_parecer_log(
                    data_hora=data_hora,
                    empresa=empresa,
                    cargo=cargo,
                    nome=nome,
                    localidade=localidade,
                    idade=idade,
                    pretensao=pretensao,
                    linkedin=linkedin,
                    resumo_profissional=resumo_prof,
                    analise_perfil=analise_prof,
                    conclusao_texto=conclusao_txt,
                    formato="PDF",
                    caminho_arquivo=caminho_final,
                    status_etapa="Em avaliação",
                    status_contratacao="Pendente",
                    motivo_decline="",
                )

                st.success(f"Parecer gerado e registrado: {caminho_final}")
                st.download_button(
                    "⬇️ Baixar PDF agora",
                    data=parecer_bytes,
                    file_name=filename,
                    mime="application/pdf",
                )

            else:  # DOCX
                parecer_bytes = build_parecer_docx_to_bytes(
                    empresa, cargo, nome, localidade, idade, pretensao,
                    resumo_prof, analise_prof, conclusao_txt,
                    linkedin
                )

                filename = nome_base + ".docx"
                caminho_final = os.path.join(output_folder, filename)

                with open(caminho_final, "wb") as f:
                    f.write(parecer_bytes)

                registrar_parecer_log(
                    data_hora=data_hora,
                    empresa=empresa,
                    cargo=cargo,
                    nome=nome,
                    localidade=localidade,
                    idade=idade,
                    pretensao=pretensao,
                    linkedin=linkedin,
                    resumo_profissional=resumo_prof,
                    analise_perfil=analise_prof,
                    conclusao_texto=conclusao_txt,
                    formato="DOCX",
                    caminho_arquivo=caminho_final,
                    status_etapa="Em avaliação",
                    status_contratacao="Pendente",
                    motivo_decline="",
                )

                st.success(f"Parecer gerado e registrado: {caminho_final}")
                st.download_button(
                    "⬇️ Baixar DOCX agora",
                    data=parecer_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        except Exception as e:
            st.error(f"Erro ao gerar o parecer: {e}")


# ==========================================================
# HISTÓRICO DE PARECERES
# ==========================================================

st.markdown("---")
st.header("📁 Histórico de pareceres registrados")

df_hist = carregar_pareceres_log()

if df_hist.empty:
    st.info("Nenhum parecer registrado ainda.")
else:
    filtro = st.text_input("Filtrar histórico por nome, empresa ou cargo")

    df_view = df_hist.copy()
    if filtro.strip():
        f = filtro.strip().lower()
        df_view = df_view[
            df_view["nome"].str.lower().str.contains(f)
            | df_view["empresa"].str.lower().str.contains(f)
            | df_view["cargo"].str.lower().str.contains(f)
        ]

    st.write(f"Total de pareceres encontrados: {len(df_view)}")

    for _, row in df_view.sort_values("data_hora", ascending=False).iterrows():
        st.markdown(
            f"""
**Data/Hora:** {row['data_hora']}  
**Empresa:** {row['empresa']}  
**Cargo:** {row['cargo']}  
**Nome:** {row['nome']}  
**Pretensão:** {row['pretensao']}  
**Formato:** {row['formato']}  
**Etapa:** {row.get('status_etapa', '')}  
**Status contratação:** {row.get('status_contratacao', '')}  
**Motivo de declínio:** {row.get('motivo_decline', '')}  
**Arquivo:** `{row['caminho_arquivo']}`
"""
        )
        arq = row["caminho_arquivo"]
        if isinstance(arq, str) and os.path.exists(arq):
            with open(arq, "rb") as f:
                dados = f.read()
            st.download_button(
                "⬇️ Baixar arquivo",
                data=dados,
                file_name=os.path.basename(arq),
                mime="application/octet-stream",
                key=f"hist_{row['data_hora']}_{row['nome']}",
            )
        else:
            st.warning("Arquivo não encontrado no caminho registrado.")
        st.markdown("---")


# ==========================================================
# PIPELINE (GRID EDITÁVEL + CARREGAR NO FORMULÁRIO)
# ==========================================================

st.header("📌 Pipeline de candidatos (gestão de etapas)")

df_pipe = carregar_pareceres_log()

if df_pipe.empty:
    st.info("Nenhum parecer registrado ainda para montar o pipeline.")
else:
    # Garante colunas de status
    for col in ["status_etapa", "status_contratacao", "motivo_decline"]:
        if col not in df_pipe.columns:
            df_pipe[col] = ""

    st.write("Edite diretamente a etapa, status de contratação e motivo de declínio:")

    etapa_opcoes = [
        "Em avaliação",
        "Triagem",
        "Entrevista",
        "Finalista",
        "Não seguiu processo",
    ]
    contratacao_opcoes = [
        "Pendente",
        "Aprovado / Contratado",
        "Reprovado",
        "Desistiu",
    ]

    cols_ordem = [
        "data_hora",
        "empresa",
        "cargo",
        "nome",
        "localidade",
        "idade",
        "pretensao",
        "status_etapa",
        "status_contratacao",
        "motivo_decline",
        "caminho_arquivo",
    ]
    df_pipe = df_pipe[[c for c in cols_ordem if c in df_pipe.columns]]

    edited_df = st.data_editor(
        df_pipe,
        hide_index=True,
        column_config={
            "status_etapa": st.column_config.SelectboxColumn(
                "Etapa",
                options=etapa_opcoes,
                help="Etapa atual do processo seletivo",
            ),
            "status_contratacao": st.column_config.SelectboxColumn(
                "Status contratação",
                options=contratacao_opcoes,
                help="Situação do candidato na vaga",
            ),
            "motivo_decline": st.column_config.TextColumn(
                "Motivo de declínio",
                help="Motivo da não contratação ou recusa",
            ),
        },
        num_rows="fixed",
        key="pipeline_editor",
    )

    if st.button("💾 Salvar alterações do pipeline"):
        try:
            edited_df.to_csv(LOG_PAR, sep=";", index=False, encoding="utf-8")
            st.success("Pipeline atualizado com sucesso!")
        except Exception as e:
            st.error(f"Erro ao salvar pipeline: {e}")

    st.markdown("### 🔁 Carregar parecer no formulário")

    df_sel = carregar_pareceres_log()
    if not df_sel.empty:
        df_sel = df_sel.sort_values("data_hora", ascending=False)
        opcoes = {
            idx: f"{row['data_hora']} - {row['nome']} - {row['cargo']} - {row['empresa']}"
            for idx, row in df_sel.iterrows()
        }

        escolha = st.selectbox(
            "Escolha um parecer para carregar no formulário:",
            options=list(opcoes.keys()),
            format_func=lambda x: opcoes[x],
        )

        if st.button("Carregar dados no formulário"):
            row = df_sel.loc[escolha]

            st.session_state["empresa"] = row.get("empresa", "")
            st.session_state["cargo"] = row.get("cargo", "")
            st.session_state["nome"] = row.get("nome", "")
            st.session_state["localidade"] = row.get("localidade", "")
            st.session_state["idade"] = str(row.get("idade", ""))
            st.session_state["pretensao"] = row.get("pretensao", "")
            st.session_state["linkedin"] = row.get("linkedin", "")
            st.session_state["resumo_profissional"] = row.get("resumo_profissional", "")
            st.session_state["analise_perfil"] = row.get("analise_perfil", "")
            st.session_state["conclusao_texto"] = row.get("conclusao_texto", "")

            st.success("Dados carregados no formulário. Role para o topo para revisar/gerar novo laudo.")
