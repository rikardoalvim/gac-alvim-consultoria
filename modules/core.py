import os
import io
import csv
import re
from datetime import datetime

import pandas as pd

from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph, Spacer, SimpleDocTemplate
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors

from PyPDF2 import PdfMerger, PdfReader
from docx import Document

# OpenAI opcional
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

# =========================
# CONFIG / PATHS
# =========================

BASE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
os.makedirs(BASE_DIR, exist_ok=True)

CV_DIR = os.path.join(BASE_DIR, "CVS")
os.makedirs(CV_DIR, exist_ok=True)

LOG_PAR = os.path.join(BASE_DIR, "pareceres_log.csv")
LOG_CAND = os.path.join(BASE_DIR, "candidatos.csv")
LOG_VAGAS = os.path.join(BASE_DIR, "vagas.csv")
LOG_VAGA_CAND = os.path.join(BASE_DIR, "vaga_candidatos.csv")
LOG_CLIENTES = os.path.join(BASE_DIR, "clientes.csv")
LOG_ACESSOS = os.path.join(BASE_DIR, "acessos.csv")

LOG_FIN_OS = os.path.join(BASE_DIR, "financeiro_os.csv")
LOG_FIN_ORC = os.path.join(BASE_DIR, "financeiro_orcamentos.csv")
LOG_FIN_NF = os.path.join(BASE_DIR, "financeiro_nf.csv")


# =========================
# OPENAI / IA
# =========================

def get_openai_client():
    if OpenAI is None:
        return None
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


def gerar_campos_via_openai(texto_base: str):
    client = get_openai_client()
    if not client:
        raise RuntimeError("OpenAI não configurado. Verifique biblioteca e OPENAI_API_KEY.")

    prompt = f"""
Você é um assistente de recrutamento.
A partir do texto abaixo (currículo + observações), gere:

[NOME]
Nome completo do candidato

[RESUMO]
Resumo Profissional (em terceira pessoa, de forma objetiva)

[ANALISE]
Análise de Perfil (postura, comunicação, senioridade, aderência à vaga, potencial)

[CONCLUSAO]
Conclusão (síntese da indicação ou não do candidato, com tom profissional)

Texto do candidato:
\"\"\"{texto_base}\"\"\"
"""

    resp = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt,
    )

    text = resp.output[0].content[0].text

    nome = resumo = analise = concl = ""
    bloco = None

    for linha in text.splitlines():
        t = linha.strip()
        if t == "[NOME]":
            bloco = "nome"
            continue
        if t == "[RESUMO]":
            bloco = "resumo"
            continue
        if t == "[ANALISE]":
            bloco = "analise"
            continue
        if t == "[CONCLUSAO]":
            bloco = "concl"
            continue

        if bloco == "nome":
            nome += linha + "\n"
        elif bloco == "resumo":
            resumo += linha + "\n"
        elif bloco == "analise":
            analise += linha + "\n"
        elif bloco == "concl":
            concl += linha + "\n"

    return nome.strip(), resumo.strip(), analise.strip(), concl.strip()


# =========================
# HELPERS GERAIS
# =========================

def _normalizar_linkedin(link):
    if not link or not link.strip():
        return None
    link = link.strip()
    if not (link.startswith("http://") or link.startswith("https://")):
        return "https://" + link
    return link


def extract_text_from_pdf(upload):
    if not upload:
        return ""
    reader = PdfReader(upload)
    text = ""
    for p in reader.pages:
        try:
            text += p.extract_text() + "\n"
        except Exception:
            pass
    return text


def montar_link_whatsapp(telefone: str) -> str:
    if not telefone:
        return ""
    digits = re.sub(r"\D", "", telefone)
    if not digits:
        return ""
    if digits.startswith("55"):
        return f"https://wa.me/{digits}"
    else:
        return f"https://wa.me/55{digits}"


# =========================
# PDFs / DOCX de PARECER
# =========================

def build_parecer_pdf_to_bytes(
    cliente, cargo, nome, localidade, idade, pretensao,
    resumo_profissional, analise_perfil, conclusao_texto,
    linkedin=""
):
    buffer = io.BytesIO()

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontSize=14,
        spaceAfter=20,
    )
    header_style = ParagraphStyle(
        "Header",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=11,
        textColor=colors.HexColor("#003366"),
    )
    section_title = ParagraphStyle(
        "SecTitle",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#003366"),
        spaceBefore=12,
        spaceAfter=4,
    )
    text_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10.5,
        alignment=TA_JUSTIFY,
        leading=15,
    )

    def to_html(txt: str) -> str:
        return txt.replace("\n", "<br/>")

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=40,
        bottomMargin=40,
        leftMargin=50,
        rightMargin=50,
    )

    content = []
    content.append(Paragraph(f"<b>Cliente:</b> {cliente}", header_style))
    content.append(Paragraph(f"<b>Cargo:</b> {cargo}", header_style))
    content.append(Spacer(1, 20))
    content.append(Paragraph("🧩 PARECER DE TRIAGEM – CANDIDATO", title_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph(
        f"<b>Nome:</b> {nome}<br/><b>Localidade:</b> {localidade}",
        text_style
    ))

    linkedin_url = _normalizar_linkedin(linkedin)
    if linkedin_url:
        content.append(Paragraph(
            f'<b>LinkedIn:</b> <link href="{linkedin_url}">{linkedin_url}</link>',
            text_style
        ))

    content.append(Paragraph("🧾 <b>Resumo Profissional</b>", section_title))
    content.append(Paragraph(to_html(resumo_profissional), text_style))

    content.append(Paragraph("💡 <b>Análise de Perfil</b>", section_title))
    content.append(Paragraph(to_html(analise_perfil), text_style))

    content.append(Paragraph("➡️ <b>Conclusão</b>", section_title))
    content.append(Paragraph(to_html(conclusao_texto), text_style))

    content.append(Paragraph("💰 <b>Informações de Remuneração</b>", section_title))
    content.append(Paragraph(
        f"Idade: {idade}<br/>Pretensão Salarial: {pretensao}",
        text_style
    ))

    doc.build(content)
    return buffer.getvalue()


def build_parecer_docx_to_bytes(
    cliente, cargo, nome, localidade, idade, pretensao,
    resumo_profissional, analise_perfil, conclusao_texto,
    linkedin=""
):
    doc = Document()

    doc.add_heading("PARECER DE TRIAGEM – CANDIDATO", level=1)
    doc.add_paragraph(f"Cliente: {cliente}")
    doc.add_paragraph(f"Cargo: {cargo}")
    doc.add_paragraph(f"Nome: {nome}")
    doc.add_paragraph(f"Localidade: {localidade}")

    linkedin_url = _normalizar_linkedin(linkedin)
    if linkedin_url:
        doc.add_paragraph(f"LinkedIn: {linkedin_url}")

    doc.add_heading("Resumo Profissional", level=2)
    for bloco in resumo_profissional.split("\n\n"):
        doc.add_paragraph(bloco)

    doc.add_heading("Análise de Perfil", level=2)
    for bloco in analise_perfil.split("\n\n"):
        doc.add_paragraph(bloco)

    doc.add_heading("Conclusão", level=2)
    for bloco in conclusao_texto.split("\n\n"):
        doc.add_paragraph(bloco)

    doc.add_heading("Informações de Remuneração", level=2)
    doc.add_paragraph(f"Idade: {idade}")
    doc.add_paragraph(f"Pretensão Salarial: {pretensao}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def merge_pdfs_bytes(parecer_bytes: bytes, resume_path: str) -> bytes:
    if not os.path.exists(resume_path):
        raise FileNotFoundError(f"Currículo não encontrado: {resume_path}")
    merger = PdfMerger()
    merger.append(io.BytesIO(parecer_bytes))
    merger.append(resume_path)

    out = io.BytesIO()
    merger.write(out)
    merger.close()
    return out.getvalue()


# =========================
# PARSE PARECER PDF (IMPORTADOR)
# =========================

def extrai_campo_simples(label: str, texto: str) -> str:
    padrao = rf"{label}\s*:\s*(.*)"
    m = re.search(padrao, texto)
    if m:
        return m.group(1).strip()
    return ""


def extrai_nome_localidade(texto: str):
    m = re.search(r"Nome\s*:\s*(.+?)Localidade\s*:\s*(.+)", texto, re.DOTALL)
    if m:
        nome = " ".join(m.group(1).strip().split())
        local = " ".join(m.group(2).strip().split())
        return nome, local
    return "", ""


def extrai_bloco(texto: str, inicio: str, fim: str | None) -> str:
    if fim:
        padrao = rf"{inicio}(.*?){fim}"
    else:
        padrao = rf"{inicio}(.*)$"
    m = re.search(padrao, texto, re.DOTALL)
    if m:
        return m.group(1).strip()
    return ""


def parse_parecer_pdf_arquivo(caminho_pdf: str) -> dict:
    resultado = {
        "cliente": "",
        "cargo": "",
        "nome": "",
        "localidade": "",
        "idade": "",
        "pretensao": "",
        "linkedin": "",
        "resumo_profissional": "",
        "analise_perfil": "",
        "conclusao_texto": "",
    }

    try:
        reader = PdfReader(caminho_pdf)
        if not reader.pages:
            return resultado
        page = reader.pages[0]
        texto = page.extract_text() or ""
    except Exception:
        return resultado

    texto_limpo = texto.replace("\r", "\n")

    cliente = extrai_campo_simples("Cliente", texto_limpo)
    cargo = extrai_campo_simples("Cargo", texto_limpo)
    nome, localidade = extrai_nome_localidade(texto_limpo)
    linkedin = extrai_campo_simples("LinkedIn", texto_limpo)

    resumo = extrai_bloco(texto_limpo, "Resumo Profissional", "Análise de Perfil")
    analise = extrai_bloco(texto_limpo, "Análise de Perfil", "Conclusão")
    conclusao = extrai_bloco(texto_limpo, "Conclusão", "Informações de Remuneração")

    info_remu = extrai_bloco(texto_limpo, "Informações de Remuneração", None)
    idade = extrai_campo_simples("Idade", info_remu)
    pret = extrai_campo_simples("Pretensão Salarial", info_remu)
    if not pret:
        pret = extrai_campo_simples("Pretensão", info_remu)

    resultado.update({
        "cliente": cliente,
        "cargo": cargo,
        "nome": nome,
        "localidade": localidade,
        "idade": idade,
        "pretensao": pret,
        "linkedin": linkedin,
        "resumo_profissional": resumo,
        "analise_perfil": analise,
        "conclusao_texto": conclusao,
    })
    return resultado


def inferir_nome_data_de_arquivo(caminho_pdf: str, nome_existente: str, data_hora_existente: str) -> tuple[str, str]:
    nome_final = nome_existente
    data_hora_final = data_hora_existente

    nome_arquivo = os.path.basename(caminho_pdf)
    base = os.path.splitext(nome_arquivo)[0]

    if not nome_final and base.lower().startswith("parecer_"):
        resto = base[len("Parecer_"):]
        partes = resto.split("_")
        if len(partes) >= 3:
            poss_data = partes[-2]
            poss_hora = partes[-1]
            nome_partes = partes[:-2]
            nome_final = " ".join(nome_partes).replace("-", " ")
        else:
            nome_final = resto.replace("_", " ").replace("-", " ")

    if not data_hora_final and base.lower().startswith("parecer_"):
        resto = base[len("Parecer_"):]
        partes = resto.split("_")
        if len(partes) >= 3:
            poss_data = partes[-2]
            poss_hora = partes[-1]
            try:
                dt = datetime.strptime(poss_data + poss_hora, "%Y%m%d%H%M")
                data_hora_final = dt.strftime("%Y-%m-%d %H:%M:%S")
            except Exception:
                pass

    if not data_hora_final:
        data_hora_final = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    return nome_final, data_hora_final


# =========================
# CLIENTES
# =========================

def carregar_clientes() -> pd.DataFrame:
    colunas = [
        "id_cliente",
        "nome_cliente",
        "razao_social",
        "cnpj",
        "cidade",
        "contato_principal",
        "telefone",
        "email",
        "observacoes",
    ]
    if not os.path.exists(LOG_CLIENTES):
        return pd.DataFrame(columns=colunas)

    df = pd.read_csv(LOG_CLIENTES, sep=";", encoding="utf-8", dtype=str).fillna("")
    changed = False
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
            changed = True
    df = df[colunas]
    if changed:
        df.to_csv(LOG_CLIENTES, sep=";", index=False, encoding="utf-8")
    return df


def registrar_cliente(nome_cliente, razao_social, cnpj, cidade, contato, telefone, email, observacoes):
    df = carregar_clientes()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_cliente"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1

    existe = os.path.exists(LOG_CLIENTES)
    with open(LOG_CLIENTES, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_cliente", "nome_cliente", "razao_social", "cnpj", "cidade",
                "contato_principal", "telefone", "email", "observacoes"
            ])
        w.writerow([
            novo_id, nome_cliente, razao_social, cnpj, cidade,
            contato, telefone, email, observacoes
        ])
    return novo_id


# =========================
# CANDIDATOS
# =========================

def carregar_candidatos() -> pd.DataFrame:
    colunas = [
        "id_candidato",
        "nome",
        "idade",
        "telefone",
        "cidade",
        "cargo_pretendido",
        "data_cadastro",
        "linkedin",
        "cv_arquivo",
    ]
    if not os.path.exists(LOG_CAND):
        return pd.DataFrame(columns=colunas)

    df = pd.read_csv(LOG_CAND, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    df = df[colunas]
    return df


def registrar_candidato(nome, idade, telefone, cidade, cargo_pretendido, data_cadastro):
    df = carregar_candidatos()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_candidato"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1

    existe = os.path.exists(LOG_CAND)
    with open(LOG_CAND, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_candidato", "nome", "idade", "telefone", "cidade",
                "cargo_pretendido", "data_cadastro", "linkedin", "cv_arquivo"
            ])
        w.writerow([
            novo_id, nome, idade, telefone, cidade,
            cargo_pretendido, data_cadastro, "", ""
        ])
    return novo_id


def get_or_create_candidato_por_nome_localidade(nome, localidade, idade, data_hora):
    df = carregar_candidatos()
    if not df.empty and nome:
        mask = df["nome"].astype(str).str.lower() == nome.lower()
        encontrados = df[mask]
        if not encontrados.empty:
            return str(encontrados.iloc[0]["id_candidato"])

    idade_str = str(idade) if idade else ""
    try:
        dt = datetime.strptime(data_hora, "%Y-%m-%d %H:%M:%S")
        data_cad = dt.strftime("%Y-%m-%d")
    except Exception:
        data_cad = datetime.now().strftime("%Y-%m-%d")

    novo_id = registrar_candidato(
        nome=nome or "Sem nome",
        idade=idade_str,
        telefone="",
        cidade=localidade or "",
        cargo_pretendido="",
        data_cadastro=data_cad,
    )
    return str(novo_id)


# =========================
# VAGAS
# =========================

def carregar_vagas() -> pd.DataFrame:
    colunas = [
        "id_vaga",
        "id_cliente",
        "nome_cliente",
        "cargo",
        "modalidade",
        "data_abertura",
        "data_fechamento",
        "status",
        "descricao_vaga",
    ]
    if not os.path.exists(LOG_VAGAS):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_VAGAS, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    df = df[colunas]
    return df


def registrar_vaga(id_cliente, nome_cliente, cargo, modalidade, data_abertura,
                   data_fechamento, status, descricao_vaga):
    df = carregar_vagas()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_vaga"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1

    existe = os.path.exists(LOG_VAGAS)
    with open(LOG_VAGAS, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_vaga", "id_cliente", "nome_cliente", "cargo", "modalidade",
                "data_abertura", "data_fechamento", "status", "descricao_vaga"
            ])
        w.writerow([
            novo_id, id_cliente, nome_cliente, cargo, modalidade,
            data_abertura, data_fechamento, status, descricao_vaga
        ])
    return novo_id


def carregar_vaga_candidatos() -> pd.DataFrame:
    colunas = ["id_vaga", "id_candidato", "data_vinculo", "observacao"]
    if not os.path.exists(LOG_VAGA_CAND):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_VAGA_CAND, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    return df[colunas]


def salvar_vaga_candidatos(df: pd.DataFrame):
    df.to_csv(LOG_VAGA_CAND, sep=";", index=False, encoding="utf-8")


# =========================
# PARECERES
# =========================

def registrar_parecer_log(
    data_hora,
    cliente,
    cargo,
    nome,
    localidade,
    idade,
    pretensao,
    linkedin,
    resumo_profissional,
    analise_perfil,
    conclusao_texto,
    formato,
    caminho_arquivo,
    id_candidato="",
    status_etapa="Em avaliação",
    status_contratacao="Pendente",
    motivo_decline="",
):
    existe = os.path.exists(LOG_PAR)
    with open(LOG_PAR, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe:
            w.writerow([
                "data_hora", "cliente", "cargo", "nome", "localidade",
                "idade", "pretensao", "linkedin",
                "resumo_profissional", "analise_perfil", "conclusao_texto",
                "formato", "caminho_arquivo",
                "id_candidato", "status_etapa", "status_contratacao", "motivo_decline"
            ])
        w.writerow([
            data_hora, cliente, cargo, nome, localidade,
            idade, pretensao, linkedin,
            resumo_profissional, analise_perfil, conclusao_texto,
            formato, caminho_arquivo,
            id_candidato, status_etapa, status_contratacao, motivo_decline
        ])


def carregar_pareceres_log() -> pd.DataFrame:
    colunas = [
        "data_hora", "cliente", "cargo", "nome", "localidade",
        "idade", "pretensao", "linkedin",
        "resumo_profissional", "analise_perfil", "conclusao_texto",
        "formato", "caminho_arquivo",
        "id_candidato", "status_etapa", "status_contratacao", "motivo_decline"
    ]
    if not os.path.exists(LOG_PAR):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_PAR, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    df = df[colunas]
    return df


# =========================
# ACESSOS (Sistemas)
# =========================

def carregar_acessos() -> pd.DataFrame:
    colunas = [
        "id_acesso",
        "id_cliente",
        "nome_cliente",
        "id_candidato",
        "nome_usuario",
        "sistema",
        "tipo_acesso",
        "data_inicio",
        "data_fim",
        "status",
        "observacoes",
    ]
    if not os.path.exists(LOG_ACESSOS):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_ACESSOS, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    df = df[colunas]
    return df


def registrar_acesso(id_cliente, nome_cliente, id_candidato, nome_usuario,
                     sistema, tipo_acesso, data_inicio, data_fim, status, observacoes):
    df = carregar_acessos()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_acesso"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1

    existe = os.path.exists(LOG_ACESSOS)
    with open(LOG_ACESSOS, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_acesso", "id_cliente", "nome_cliente", "id_candidato",
                "nome_usuario", "sistema", "tipo_acesso",
                "data_inicio", "data_fim", "status", "observacoes"
            ])
        w.writerow([
            novo_id, id_cliente, nome_cliente, id_candidato,
            nome_usuario, sistema, tipo_acesso,
            data_inicio, data_fim, status, observacoes
        ])
    return novo_id


# =========================
# FINANCEIRO
# =========================

def carregar_fin_os() -> pd.DataFrame:
    colunas = [
        "id_os", "id_cliente", "nome_cliente",
        "descricao", "tipo_servico", "data_emissao",
        "data_execucao", "valor", "status", "observacoes"
    ]
    if not os.path.exists(LOG_FIN_OS):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_FIN_OS, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    return df[colunas]


def registrar_fin_os(id_cliente, nome_cliente, descricao, tipo_servico,
                     data_emissao, data_execucao, valor, status, observacoes):
    df = carregar_fin_os()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_os"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1
    existe = os.path.exists(LOG_FIN_OS)
    with open(LOG_FIN_OS, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_os", "id_cliente", "nome_cliente",
                "descricao", "tipo_servico", "data_emissao",
                "data_execucao", "valor", "status", "observacoes"
            ])
        w.writerow([
            novo_id, id_cliente, nome_cliente,
            descricao, tipo_servico, data_emissao,
            data_execucao, valor, status, observacoes
        ])
    return novo_id


def carregar_fin_orc() -> pd.DataFrame:
    colunas = [
        "id_orc", "id_cliente", "nome_cliente",
        "descricao", "data_emissao", "validade",
        "valor", "status", "observacoes"
    ]
    if not os.path.exists(LOG_FIN_ORC):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_FIN_ORC, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    return df[colunas]


def registrar_fin_orc(id_cliente, nome_cliente, descricao, data_emissao,
                      validade, valor, status, observacoes):
    df = carregar_fin_orc()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_orc"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1
    existe = os.path.exists(LOG_FIN_ORC)
    with open(LOG_FIN_ORC, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_orc", "id_cliente", "nome_cliente",
                "descricao", "data_emissao", "validade",
                "valor", "status", "observacoes"
            ])
        w.writerow([
            novo_id, id_cliente, nome_cliente,
            descricao, data_emissao, validade,
            valor, status, observacoes
        ])
    return novo_id


def carregar_fin_nf() -> pd.DataFrame:
    colunas = [
        "id_nf", "id_cliente", "nome_cliente",
        "numero_nf", "data_emissao",
        "valor", "descricao", "observacoes"
    ]
    if not os.path.exists(LOG_FIN_NF):
        return pd.DataFrame(columns=colunas)
    df = pd.read_csv(LOG_FIN_NF, sep=";", encoding="utf-8", dtype=str).fillna("")
    for c in colunas:
        if c not in df.columns:
            df[c] = ""
    return df[colunas]


def registrar_fin_nf(id_cliente, nome_cliente, numero_nf, data_emissao,
                     valor, descricao, observacoes):
    df = carregar_fin_nf()
    if df.empty:
        novo_id = 1
    else:
        try:
            novo_id = int(pd.to_numeric(df["id_nf"], errors="coerce").max() or 0) + 1
        except Exception:
            novo_id = 1
    existe = os.path.exists(LOG_FIN_NF)
    with open(LOG_FIN_NF, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        if not existe or df.empty:
            w.writerow([
                "id_nf", "id_cliente", "nome_cliente",
                "numero_nf", "data_emissao",
                "valor", "descricao", "observacoes"
            ])
        w.writerow([
            novo_id, id_cliente, nome_cliente,
            numero_nf, data_emissao,
            valor, descricao, observacoes
        ])
    return novo_id
